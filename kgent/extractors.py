"""Text extraction for binary document formats: PDF, Word, and email.

This module keeps ingest.py free of format specific code. Email parsing uses
the Python standard library, while PDF and Word need the optional ``docs``
extras (``pymupdf`` and ``python-docx``).
"""
from __future__ import annotations

import email
import html
import io
import mailbox
import re
from collections.abc import Iterator
from dataclasses import dataclass, field
from email.header import decode_header
from email.message import Message
from pathlib import Path

_HTML_TAG_RE = re.compile(r"<[^>]+>")
_TEXT_ATTACHMENT_SUFFIXES = {".txt", ".md", ".csv", ".log"}


@dataclass(frozen=True)
class EmailContent:
    """The readable text of one email and its raw attachments."""

    text: str
    attachments: list[tuple[str, bytes]] = field(default_factory=list)


def extract_pdf(data: bytes) -> str:
    """Return the plain text of a PDF document held in memory."""
    import fitz

    parts: list[str] = []
    with fitz.open(stream=data, filetype="pdf") as doc:
        for page in doc:
            parts.append(page.get_text())
    return "\n".join(parts).strip()


def extract_docx(data: bytes) -> str:
    """Return the plain text of a Word .docx document held in memory."""
    from docx import Document as DocxDocument

    doc = DocxDocument(io.BytesIO(data))
    parts: list[str] = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts).strip()


def extract_attachment(filename: str, data: bytes) -> str | None:
    """Extract text from an email attachment, or None for unsupported types."""
    suffix = Path(filename).suffix.lower()
    if suffix == ".pdf":
        return extract_pdf(data)
    if suffix == ".docx":
        return extract_docx(data)
    if suffix in _TEXT_ATTACHMENT_SUFFIXES:
        return data.decode("utf-8", errors="replace").strip()
    return None


def parse_email(data: bytes) -> EmailContent:
    """Parse a single .eml message into readable text and attachments."""
    return _email_content(email.message_from_bytes(data))


def iter_mbox(path: Path) -> Iterator[EmailContent]:
    """Yield one EmailContent per message in a Thunderbird style mbox file."""
    box = mailbox.mbox(str(path))
    try:
        for message in box:
            yield _email_content(message)
    finally:
        box.close()


def _email_content(msg: Message) -> EmailContent:
    headers: list[str] = []
    for key in ("Subject", "From", "To", "Date"):
        value = msg.get(key)
        if value:
            headers.append(f"{key}: {_decode_header(value)}")

    body_plain: list[str] = []
    body_html: list[str] = []
    attachments: list[tuple[str, bytes]] = []
    for part in msg.walk():
        if part.is_multipart():
            continue
        filename = part.get_filename()
        if filename:
            payload = part.get_payload(decode=True)
            if payload:
                attachments.append((_decode_header(filename), payload))
            continue
        content_type = part.get_content_type()
        if content_type == "text/plain":
            body_plain.append(_decode_part(part))
        elif content_type == "text/html":
            body_html.append(_decode_part(part))

    body = "\n".join(t for t in body_plain if t).strip()
    if not body and body_html:
        body = _strip_html("\n".join(body_html)).strip()

    text = "\n".join(headers)
    if body:
        text = f"{text}\n\n{body}" if text else body
    return EmailContent(text=text.strip(), attachments=attachments)


def _decode_part(part: Message) -> str:
    payload = part.get_payload(decode=True)
    if not payload:
        return ""
    charset = part.get_content_charset() or "utf-8"
    try:
        return payload.decode(charset, errors="replace")
    except (LookupError, ValueError):
        return payload.decode("utf-8", errors="replace")


def _decode_header(value: str) -> str:
    parts: list[str] = []
    for chunk, encoding in decode_header(value):
        if isinstance(chunk, bytes):
            parts.append(chunk.decode(encoding or "utf-8", errors="replace"))
        else:
            parts.append(chunk)
    return " ".join(p.strip() for p in parts if p and p.strip())


def _strip_html(raw: str) -> str:
    text = _HTML_TAG_RE.sub(" ", raw)
    text = html.unescape(text)
    return re.sub(r"[ \t]+", " ", text)
